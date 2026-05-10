from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

document = Document()

# Title
head = document.add_heading('Project Proposal: AquaClear Underwater Image Enhancement System', 0)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = document.add_paragraph()
p.add_run('Course: ').bold = True
p.add_run('CENG 455 Digital Image Processing\n')
p.add_run('Institution: ').bold = True
p.add_run('Istanbul Arel University, 2026\n')
p.add_run('Student: ').bold = True
p.add_run('Khaled Noaman')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_heading('1. Introduction and Background', level=1)
p = document.add_paragraph(
    "Underwater photography and videography suffer from severe degradation compared to terrestrial imaging. "
    "Due to the physical properties of water, light is exponentially attenuated and scattered as it travels "
    "from the surface to the object, and back to the camera lens. This results in three primary issues:"
)
document.add_paragraph("Color Cast (Attenuation): Water absorbs wavelengths differently. Red light (650nm) is absorbed first, leaving a dominant blue/green tint.", style='List Bullet')
document.add_paragraph("Haze and Fog (Backscattering): Suspended particles scatter light, creating a thick fog.", style='List Bullet')
document.add_paragraph("Loss of Detail (Forward Scattering): Light reflecting off the object is scattered slightly, blurring edges.", style='List Bullet')

document.add_heading('2. Problem Statement', level=1)
p = document.add_paragraph(
    "Standard image enhancement techniques (such as global histogram equalization or basic brightness/contrast "
    "adjustments) are insufficient for underwater images. They often wash out scenes and amplify noise. "
    "A specialized, multi-stage image processing pipeline is required to reverse specific aquatic degradations."
)

document.add_heading('3. Proposed Solution', level=1)
p = document.add_paragraph(
    "AquaClear is a web-based software application that implements a purely classical Image Processing (IP) "
    "pipeline to restore underwater images. It avoids Neural Networks in favor of highly interpretable algorithms. "
    "The pipeline consists of:"
)
document.add_paragraph("Non-Linear Noise Reduction: Bilateral Filtering to remove noise while preserving edges.", style='List Bullet')
document.add_paragraph("Color Correction: Gray World Theorem to globally estimate color deficits and neutralize tints.", style='List Bullet')
document.add_paragraph("Dehazing (Dark Channel Prior): Estimating the transmission map to subtract backscattered haze.", style='List Bullet')
document.add_paragraph("Adaptive Contrast: Contrast Limited Adaptive Histogram Equalization (CLAHE) on the LAB Luminance channel.", style='List Bullet')
document.add_paragraph("Frequency Sharpening: Unsharp Masking to counteract forward-scatter blur.", style='List Bullet')

document.add_heading('4. System Architecture', level=1)
document.add_paragraph("Frontend: Responsive HTML/CSS/JS with Dark Mode and visual metrics.", style='List Bullet')
document.add_paragraph("Backend: Node.js (Express) server handling uploads and API routing.", style='List Bullet')
document.add_paragraph("IP Engine: Python 3 with OpenCV (cv2) executing matrix operations.", style='List Bullet')

document.add_heading('5. Key Feature: Auto-Analysis', level=1)
p = document.add_paragraph(
    "When an image is uploaded, the Python backend performs a preliminary diagnostic pass to check global variance (blur) "
    "and dark channel mean (haze). Based on these heuristics, the system automatically configures the optimal "
    "parameters for the enhancement pipeline."
)

document.add_heading('6. Expected Outcomes', level=1)
document.add_paragraph("Functional Web Application handling real-time image enhancement.", style='List Number')
document.add_paragraph("Source Code Repository with clean, documented modular code.", style='List Number')
document.add_paragraph("Project Presentation detailing IP mechanics.", style='List Number')
document.add_paragraph("Final Report evaluating quantitative results (PSNR/SSIM).", style='List Number')

out_path = os.path.join(os.path.dirname(__file__), 'AquaClear_Project_Proposal.docx')
document.save(out_path)
print(f"Proposal saved to {out_path}")
